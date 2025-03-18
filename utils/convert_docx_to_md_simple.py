from docx import Document
from .helper import clean_markdown
def process_paragraph_text(paragraph):
    """Xử lý text trong paragraph với định dạng bold/italic"""
    runs = paragraph.runs
    if not runs:
        return ""
        
    # Gộp các run có cùng định dạng
    formatted_parts = []
    current_text = ""
    current_format = None
    
    for run in runs:
        # Chỉ xét bold và italic
        format_type = (run.bold, run.italic)
        
        if format_type != current_format and current_text:
            if current_format:
                is_bold, is_italic = current_format
                if is_bold and is_italic:
                    current_text = f"***{current_text}***"
                elif is_bold:
                    current_text = f"**{current_text}**"
                elif is_italic:
                    current_text = f"_{current_text}_"
            formatted_parts.append(current_text)
            current_text = ""
            
        current_text += run.text
        current_format = format_type
    
    # Xử lý phần text cuối cùng
    if current_text:
        if current_format:
            is_bold, is_italic = current_format
            if is_bold and is_italic:
                current_text = f"***{current_text}***"
            elif is_bold:
                current_text = f"**{current_text}**"
            elif is_italic:
                current_text = f"_{current_text}_"
        formatted_parts.append(current_text)
    
    return "".join(formatted_parts).strip()

def get_cell_alignment(cell):
    """Xác định căn chỉnh của cell dựa vào paragraph alignment"""
    for paragraph in cell.paragraphs:
        if paragraph.alignment:
            if paragraph.alignment == 1:
                return "center"
            elif paragraph.alignment == 2:
                return "right"
    return "left"

def convert_word_to_markdown_simple(doc_path):
    doc = Document(doc_path)
    md_text = ""

    for element in doc.element.body:
        if element.tag.endswith("p"):  # Đoạn văn bình thường
            paragraph = next(p for p in doc.paragraphs if p._element == element)
            text = process_paragraph_text(paragraph)
            if text:
                md_text += text + "\n\n"

        elif element.tag.endswith("tbl"):  # Bảng
            table = next(t for t in doc.tables if t._element == element)
            
            # Xử lý header và xác định alignment cho mỗi cột
            headers = []
            alignments = []
            for cell in table.rows[0].cells:
                cell_text = []
                for para in cell.paragraphs:
                    text = process_paragraph_text(para)
                    if text:
                        cell_text.append(text)
                headers.append(" ".join(cell_text) if cell_text else " ")
                alignments.append(get_cell_alignment(cell))
            
            # Tạo bảng markdown với alignment
            md_text += "| " + " | ".join(headers) + " |\n"
            
            # Tạo delimiter row với alignment indicators
            delimiter_row = []
            for align in alignments:
                if align == "center":
                    delimiter_row.append(":---:")
                elif align == "right":
                    delimiter_row.append("---:")
                else:  # left
                    delimiter_row.append(":---")
            md_text += "|" + "|".join(delimiter_row) + "|\n"

            # Xử lý nội dung bảng
            for row in table.rows[1:]:
                row_data = []
                for cell in row.cells:
                    cell_text = []
                    for para in cell.paragraphs:
                        text = process_paragraph_text(para)
                        if text:
                            cell_text.append(text)
                    row_data.append("<br>".join(cell_text) if cell_text else " ")
                
                md_text += "| " + " | ".join(row_data) + " |\n"
            
            md_text += "\n"

    return clean_markdown(md_text)